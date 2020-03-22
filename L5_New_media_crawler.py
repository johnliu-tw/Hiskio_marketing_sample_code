from selenium import webdriver
from selenium.webdriver.chrome.options import Options 
from bs4 import BeautifulSoup
import time 
import os
import pymysql.cursors
import urllib.request
from datetime import date
today = str(date.today())

from docx import Document


connection = pymysql.connect(host='localhost', 
                             user='root', 
                             password='password', 
                             db='new_media', 
                             cursorclass=pymysql.cursors.DictCursor)

options = Options()
options.add_argument('--headless')
options.experimental_options["prefs"] = {'profile.default_content_settings' : {"images": 2}, 
                                        'profile.managed_default_content_settings' :  {"images": 2}}

document = Document()
document.add_heading(today + '新媒體文章資料爬取', 0)
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '標題'
hdr_cells[1].text = '分享數'
hdr_cells[2].text = 'Hashtag'
hdr_cells[3].text = '網站'

driver = webdriver.Chrome(os.getcwd() + '/chromedriver', chrome_options=options)
try:
    with connection.cursor() as cursor:
        for i in range(1, 3):
            driver.get('https://www.inside.com.tw/?page='+ str(i) )
            sourceCode = BeautifulSoup(driver.page_source)
            article_box = sourceCode.select('div.post_list-list_style')[0]
            articles = article_box.select('div.post_list_item')
            for article in articles:
                title = article.select('h3.post_title')[0].text
                date = article.select('li.post_date')[0].text.strip().replace('/', '-')
                tags = article.select('a.hero_slide_tag')
                tags_string = ''
                for tag in tags:
                    tags_string += tag.text + ', '
                if today == date:
                    print(title)
                    print(date)
                    print(tags_string)

                    row_cells = table.add_row().cells
                    row_cells[0].text = title
                    row_cells[1].text = ''
                    row_cells[2].text = tags_string
                    row_cells[3].text = 'Inside'
                    
                    sql = '''
                    INSERT INTO `new_media`.`articles` (`title`, `date`, `tags`, `brand`)
                    VALUES('{}', '{}', '{}', '{}')
                    '''.format(title, date, tags_string, 'inside')  
                    cursor.execute(sql)
                    connection.commit()
            
            driver.get('https://technews.tw/page/'+ str(i) +'/')
            sourceCode = BeautifulSoup(driver.page_source)
            article_box = sourceCode.select('div#content')[0]
            articles = article_box.select('header.entry-header')
            for article in articles:
                title = article.select('h1.entry-title')[0].text
                date = article.select('span.body')[1].text.strip().replace(' 年 ', '-').replace(' 月 ', '-').replace(' 日 ', '-')
                date = date[0:10]
                tags = article.select('span.body')[2].select('a')
                iframe = article.select('iframe')[1]
                response = urllib.request.urlopen(iframe.attrs['src'])
                iframe_soup = BeautifulSoup(response)
                share = iframe_soup.select('span#u_0_2')[0].text
                tags_string = ''
                for tag in tags:
                    tags_string += tag.text + ', '
                if today == date:
                    print(title)
                    print(date)
                    print(tags_string)
                    print(share)

                    row_cells = table.add_row().cells
                    row_cells[0].text = title
                    row_cells[1].text = share
                    row_cells[2].text = tags_string
                    row_cells[3].text = 'TechNews'
                    sql = '''
                    INSERT INTO `new_media`.`articles` (`title`, `date`, `tags`, `share`, `brand`)
                    VALUES('{}', '{}', '{}', '{}', '{}')
                    '''.format(title, date, tags_string, share , 'technews')  
                    cursor.execute(sql)
                    connection.commit()
            
            
        driver.get('https://buzzorange.com/techorange/')
        for i in range(1, 4):
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
        sourceCode = BeautifulSoup(driver.page_source)
        article_box = sourceCode.select('main#main')[0]
        articles = article_box.select('article')
        for article in articles:
            title = article.select('h4.entry-title')[0].text
            date = article.select('time.entry-date')[0].text.strip().replace('/', '-')
            share = article.select('span.shareCount')[0].text
            if share.find('K') != -1:
                share = float(share.split(' ')[0]) * 1000
            else:
                share = share.split(' ')[0]
                
            if today == date:
                print(title)
                print(date)
                print(share)
                row_cells = table.add_row().cells
                row_cells[0].text = title
                row_cells[1].text = share
                row_cells[2].text = ''
                row_cells[3].text = '科技報橘'
                sql = '''
                INSERT INTO `new_media`.`articles` (`title`, `date`, `share`, `brand`)
                VALUES('{}', '{}', '{}', '{}')
                '''.format(title, date, share , 'techorange')  
                cursor.execute(sql)
                connection.commit()

        driver.get('https://www.storm.mg/articles')
        sourceCode = BeautifulSoup(driver.page_source)
        article_box = sourceCode.select('div.category_cards_wrapper')[0]
        articles = article_box.select('div.category_card')
        for article in articles:
            title = article.select('.card_title')[0].text
            date = article.select('span.info_time')[0].text
            tags = article.select('a.card_tag')
            tags_string = ''
            for tag in tags:
                tags_string += tag.text + ','

            print(title)
            print(date)
            print(tags_string)

                
            
#     document.save(today + '新媒體文章爬取.docx')
    connection.close()
    driver.close()
except Exception as e:
    print(e)
    connection.close()
    driver.close()
